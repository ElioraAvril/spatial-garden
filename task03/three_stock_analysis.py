# -*- coding: utf-8 -*-
"""
三股双均线策略综合分析
========================================
1) 金发科技(600143.SH) + 宁德时代(300750.SZ) + 比亚迪(002594.SZ)
2) 统一 MA5/MA15 金叉死叉策略
3) 整合到同一交互式 HTML 页面（可切换选择股票）
4) 输出各股的 Word 报告
5) 双均线策略适用场景总结
"""

import os, sys, base64, warnings, re, tempfile, json
warnings.filterwarnings('ignore')
os.environ['MPLCONFIGDIR'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.matplotlib_cache')
from datetime import datetime, timedelta
from io import BytesIO

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

# ===== 字体 =====
simsun_path = 'C:/Windows/Fonts/simsun.ttc'
fm.fontManager.addfont(simsun_path)
plt.rcParams['font.family'] = 'SimSun'
plt.rcParams['axes.unicode_minus'] = False

# ===== 输出目录 =====
OUT_DIR = 'D:/项目/北大BA/在线实习/update'
os.makedirs(OUT_DIR, exist_ok=True)

# ===== 参数 =====
STOCKS = [
    {'code': '600143.SH', 'name': '金发科技'},
    {'code': '300750.SZ', 'name': '宁德时代'},
    {'code': '002594.SZ', 'name': '比亚迪'},
]
SHORT_WIN = 5
LONG_WIN = 15
INITIAL_CAPITAL = 1000000
RISK_FREE_RATE = 0.025

print('=' * 60)
print('三股双均线策略综合分析 (MA{}/MA{})'.format(SHORT_WIN, LONG_WIN))
print('=' * 60)

# ============================================================
# 核心分析函数
# ============================================================
import tushare as ts
pro = ts.pro_api()

def analyze_stock(stock_info):
    """对单只股票执行完整分析"""
    code, name = stock_info['code'], stock_info['name']
    print(f'\n【{name}】分析中...')

    end_date = datetime.now().strftime('%Y%m%d')
    start_date = (datetime.now() - timedelta(days=400)).strftime('%Y%m%d')

    # 获取日线
    df = pro.daily(ts_code=code, start_date=start_date, end_date=end_date)
    if df is None or df.empty:
        raise Exception(f'{name}({code}) 无数据')
    df.sort_values('trade_date', inplace=True)
    df.reset_index(drop=True, inplace=True)

    # 前复权
    df_adj = pro.daily(ts_code=code, start_date=start_date, end_date=end_date, adj='qfq')
    df_adj.sort_values('trade_date', inplace=True)
    df_adj.reset_index(drop=True, inplace=True)
    df['open_adj'] = df_adj['open']
    df['close_adj'] = df_adj['close']
    df['high_adj'] = df_adj['high']
    df['low_adj'] = df_adj['low']
    df['date'] = pd.to_datetime(df['trade_date'])

    # 近一年
    one_year = df['date'].max() - pd.DateOffset(days=365)
    dp = df[df['date'] >= one_year].copy()
    dp.reset_index(drop=True, inplace=True)
    print(f'  近一年数据: {len(dp)} 条 ({dp.trade_date.iloc[0]} ~ {dp.trade_date.iloc[-1]})')

    # 均线
    dp['MA_short'] = dp['close_adj'].rolling(SHORT_WIN).mean()
    dp['MA_long']  = dp['close_adj'].rolling(LONG_WIN).mean()

    # 信号
    dp['signal'] = 0
    dp.loc[(dp['MA_short'] > dp['MA_long']) & (dp['MA_short'].shift(1) <= dp['MA_long'].shift(1)), 'signal'] = 1
    dp.loc[(dp['MA_short'] < dp['MA_long']) & (dp['MA_short'].shift(1) >= dp['MA_long'].shift(1)), 'signal'] = -1

    buy_sig  = dp[dp['signal'] == 1].copy()
    sell_sig = dp[dp['signal'] == -1].copy()
    print(f'  买入信号: {len(buy_sig)} 次, 卖出信号: {len(sell_sig)} 次')

    # 回测
    pos = 0; cash = INITIAL_CAPITAL; nav_list = []; trades = []
    for i in range(len(dp)):
        row = dp.iloc[i]; pr = row['close_adj']
        if row['signal'] == 1 and cash > 0:
            sh = int(cash / pr / 100) * 100
            if sh > 0:
                cst = sh * pr; cash -= cst; pos += sh
                trades.append({'date': row['trade_date'], 'type': '买入', 'price': pr, 'shares': sh, 'amount': cst, 'cash': cash})
        elif row['signal'] == -1 and pos > 0:
            amt = pos * pr; cash += amt; pos = 0
            trades.append({'date': row['trade_date'], 'type': '卖出', 'price': pr, 'shares': 0, 'amount': amt, 'cash': cash})
        nav_list.append(cash + pos * pr)
    dp['nav'] = nav_list

    nav_arr = np.array(nav_list)
    final_val = nav_arr[-1]
    total_ret = (final_val - INITIAL_CAPITAL) / INITIAL_CAPITAL * 100
    days = (dp['date'].iloc[-1] - dp['date'].iloc[0]).days
    ann_ret = (final_val / INITIAL_CAPITAL) ** (365 / max(days, 1)) - 1
    ann_ret_pct = ann_ret * 100

    peak = np.maximum.accumulate(nav_arr)
    dd = (nav_arr - peak) / peak
    mdd = dd.min() * 100

    dr = pd.Series(nav_arr).pct_change().dropna()
    sharpe = ((dr.mean() * 252) - RISK_FREE_RATE) / (dr.std() * np.sqrt(252)) if dr.std() > 0 else 0

    wins = 0; losses = 0; tr = []
    for i in range(0, len(trades)-1, 2):
        if i+1 < len(trades) and trades[i]['type']=='买入' and trades[i+1]['type']=='卖出':
            r = (trades[i+1]['amount'] - trades[i]['amount']) / trades[i]['amount'] * 100
            tr.append(r)
            if r > 0: wins += 1
            else: losses += 1
    wt = wins + losses
    winr = wins / wt * 100 if wt > 0 else 0
    avg_tr = np.mean(tr) if tr else 0

    bh_ret = (dp['close_adj'].iloc[-1] - dp['close_adj'].iloc[0]) / dp['close_adj'].iloc[0] * 100
    ex_ret = total_ret - bh_ret
    calmar = ann_ret_pct / abs(mdd) if mdd != 0 else 0

    # 描述
    first_c = dp['close_adj'].iloc[0]
    last_c  = dp['close_adj'].iloc[-1]
    pchg = (last_c - first_c) / first_c * 100
    if pchg > 10: trend = '单边上涨'
    elif pchg > 0: trend = '震荡偏强'
    elif pchg > -10: trend = '震荡偏弱'
    else: trend = '单边下跌'

    print(f'  累计收益: {total_ret:+.2f}% | 年化: {ann_ret_pct:+.2f}% | MDD: {mdd:.2f}%')
    print(f'  夏普: {sharpe:.2f} | 胜率: {winr:.1f}% | 买入持有: {bh_ret:+.2f}%')

    metrics = {
        'code': code, 'name': name,
        'initial': INITIAL_CAPITAL, 'final': final_val,
        'total_return': total_ret, 'annual_return': ann_ret_pct,
        'mdd': mdd, 'sharpe': sharpe, 'calmar': calmar,
        'trades_total': len(trades),
        'wins': wins, 'losses': losses, 'win_rate': winr,
        'avg_trade_return': avg_tr, 'days': days,
        'buy_hold_return': bh_ret, 'excess_return': ex_ret,
        'volatility': dr.std() * np.sqrt(252) * 100,
        'first_close': first_c, 'last_close': last_c,
        'price_change': pchg, 'trend': trend,
        'trade_log': trades,
        'buy_signals': buy_sig, 'sell_signals': sell_sig,
        'df_plot': dp,
        'nav_arr': nav_arr,
        'first_date': dp['trade_date'].iloc[0],
        'last_date': dp['trade_date'].iloc[-1],
        'total_rows': len(dp),
    }

    # 生成回测图表
    fig = plt.figure(figsize=(16, 12), facecolor='white')
    ax1 = fig.add_axes([0.07, 0.38, 0.88, 0.54])
    ax1.set_facecolor('#ffffff')
    ax1.plot(dp['date'], dp['close_adj'], color='#2c3e50', lw=1.0, alpha=0.6, label='复权收盘价', zorder=1)
    ax1.plot(dp['date'], dp['MA_short'], color='#e74c3c', lw=2.2, label=f'MA{SHORT_WIN}(短期)', zorder=2)
    ax1.plot(dp['date'], dp['MA_long'], color='#2980b9', lw=2.2, label=f'MA{LONG_WIN}(长期)', zorder=2)
    ax1.fill_between(dp['date'], dp['MA_short'], dp['MA_long'],
                     where=(dp['MA_short'] >= dp['MA_long']), facecolor='#e74c3c', alpha=0.06, zorder=0)
    ax1.fill_between(dp['date'], dp['MA_short'], dp['MA_long'],
                     where=(dp['MA_short'] < dp['MA_long']), facecolor='#2980b9', alpha=0.06, zorder=0)
    ax1.scatter(buy_sig['date'], buy_sig['close_adj'], marker='^', s=220,
                color='#e74c3c', edgecolors='#c0392b', lw=1.5, zorder=5, label='买入(金叉)')
    ax1.scatter(sell_sig['date'], sell_sig['close_adj'], marker='v', s=220,
                color='#27ae60', edgecolors='#1e8449', lw=1.5, zorder=5, label='卖出(死叉)')
    for cnt, (idx, row) in enumerate(buy_sig.iterrows()):
        if cnt in [0, len(buy_sig)-1]:
            ax1.annotate(' 买入', (row['date'], row['close_adj']), xytext=(8,12),
                        textcoords='offset points', fontsize=9, color='#c0392b', fontweight='bold',
                        arrowprops=dict(arrowstyle='->', color='#c0392b', lw=0.8))
    for cnt, (idx, row) in enumerate(sell_sig.iterrows()):
        if cnt in [0, len(sell_sig)-1]:
            ax1.annotate(' 卖出', (row['date'], row['close_adj']), xytext=(8,-18),
                        textcoords='offset points', fontsize=9, color='#1e8449', fontweight='bold',
                        arrowprops=dict(arrowstyle='->', color='#1e8449', lw=0.8))
    ax1.set_title(f'{name}({code}) 双均线策略回测  MA{SHORT_WIN}/MA{LONG_WIN}',
                  fontsize=16, fontweight='bold', pad=15, color='#2c3e50')
    ax1.set_ylabel('价格(元)', fontsize=12)
    ax1.legend(loc='upper left', fontsize=9, framealpha=0.9, ncol=2)
    ax1.grid(True, alpha=0.15, linestyle='--')
    ax1.spines['top'].set_visible(False); ax1.spines['right'].set_visible(False)

    ax2 = fig.add_axes([0.07, 0.24, 0.88, 0.12])
    ax2.set_facecolor('#ffffff')
    cv = ['#e74c3c' if c < o else '#27ae60' for c, o in zip(dp['close_adj'], dp['open_adj'])]
    ax2.bar(dp['date'], dp['vol'], color=cv, alpha=0.4, width=0.8, label='成交量')
    ax2.set_ylabel('成交量(手)', fontsize=10)
    ax2.legend(loc='upper left', fontsize=8)
    ax2.grid(True, alpha=0.12, linestyle='--')
    ax2.spines['top'].set_visible(False); ax2.spines['right'].set_visible(False)
    ax2.set_xticklabels([])

    ax3 = fig.add_axes([0.07, 0.07, 0.88, 0.15])
    ax3.set_facecolor('#ffffff')
    ax3.plot(dp['date'], nav_list, color='#8e44ad', lw=1.8, label='策略净值', zorder=2)
    ax3.fill_between(dp['date'], INITIAL_CAPITAL, nav_list,
                     where=(np.array(nav_list) >= INITIAL_CAPITAL), facecolor='#e74c3c', alpha=0.12)
    ax3.fill_between(dp['date'], INITIAL_CAPITAL, nav_list,
                     where=(np.array(nav_list) < INITIAL_CAPITAL), facecolor='#27ae60', alpha=0.12)
    ax3.axhline(y=INITIAL_CAPITAL, color='#7f8c8d', ls='--', lw=0.8, alpha=0.6)
    ax3.set_ylabel('净值(万元)', fontsize=10)
    ax3.set_xlabel('交易日期', fontsize=11)
    ax3.legend(loc='upper left', fontsize=8)
    ax3.grid(True, alpha=0.12, linestyle='--')
    ax3.spines['top'].set_visible(False); ax3.spines['right'].set_visible(False)
    plt.xticks(rotation=30)

    chart_file = os.path.join(OUT_DIR, f'{name}_双均线策略回测图.png')
    fig.savefig(chart_file, dpi=200, bbox_inches='tight', facecolor='white')
    buf = BytesIO()
    fig.savefig(buf, dpi=180, bbox_inches='tight', facecolor='white', edgecolor='none')
    buf.seek(0)
    chart_b64 = base64.b64encode(buf.read()).decode('utf-8')
    plt.close(fig)

    metrics['chart_path'] = chart_file
    metrics['chart_b64'] = chart_b64
    return metrics


# ============================================================
# 逐一分析
# ============================================================
results = []
for s in STOCKS:
    m = analyze_stock(s)
    results.append(m)
    print(f'  [{m["name"]}] 回测完成')

print('\n' + '=' * 60)
print('三股综合分析完成')
print('=' * 60)

# ============================================================
# 生成交互式 HTML 页面
# ============================================================
print('\n生成交互式 HTML 页面...')

# 构建各股票的 JS 数据
stock_js_list = []

for m in results:
    name = m['name']
    code = m['code']

    # 交易记录HTML
    tr_html = ''
    for i, t in enumerate(m['trade_log'], 1):
        cls = 'signal-buy' if t['type'] == '买入' else 'signal-sell'
        tr_html += f'<tr><td>{i}</td><td>{t["date"]}</td><td class="{cls}">{t["type"]}</td><td>{t["price"]:.2f}</td><td>{int(t["shares"]):,}</td><td>{t["amount"]:,.0f}</td><td>{t["cash"]:,.0f}</td></tr>\n'

    # 买卖信号
    buy_dates = [r['trade_date'] for _, r in m['buy_signals'].iterrows()]
    sell_dates = [r['trade_date'] for _, r in m['sell_signals'].iterrows()]

    # 信号详情
    buy_detail = ''.join(f'<li>{d} 价格 {m["buy_signals"].iloc[i]["close_adj"]:.2f}元 &#x2191;</li>' for i, d in enumerate(buy_dates))
    sell_detail = ''.join(f'<li>{d} 价格 {m["sell_signals"].iloc[i]["close_adj"]:.2f}元 &#x2193;</li>' for i, d in enumerate(sell_dates))

    # 解读文本
    if m['total_return'] > 0:
        perf_icon = '✅ 策略盈利'
        perf_color = '#27ae60'
    else:
        perf_icon = '❌ 策略亏损'
        perf_color = '#e74c3c'

    if m['excess_return'] > 0:
        vs_text = f'策略跑赢买入持有 <span style="color:#27ae60;">+{m["excess_return"]:.2f}%</span>'
    else:
        vs_text = f'策略跑输买入持有 <span style="color:#e74c3c;">{m["excess_return"]:.2f}%</span>'

    if m['sharpe'] > 1:
        sp_text = '夏普比率表现较好，风险调整后收益理想'
    elif m['sharpe'] > 0:
        sp_text = '夏普比率表现一般'
    else:
        sp_text = '夏普比率为负，风险调整后收益需改进'

    # 交易频率
    freq = len(m['trade_log']) / max(m['days']/365, 0.1)

    # 信号频率
    total_signals = len(buy_dates) + len(sell_dates)
    signal_interval = m['days'] // total_signals if total_signals > 0 else 0

    # 构建 HTML 块
    stock_html = f'''
    <!-- ===== {name} ({code}) ===== -->
    <div class="stock-content" id="stock-{code.replace('.','')}">
        <div class="stock-header">
            <h2>{name}<span class="stock-code">{code}</span></h2>
            <div class="stock-summary">
                <span class="summary-item" style="background:{perf_color}20;color:{perf_color};border:1px solid {perf_color}30;">{perf_icon}</span>
                <span class="summary-item">回测周期: {m['first_date']} ~ {m['last_date']}</span>
                <span class="summary-item">数据量: {m['total_rows']}个交易日</span>
                <span class="summary-item">均线参数: MA{SHORT_WIN}/MA{LONG_WIN}</span>
                <span class="summary-item">{m['trend']}</span>
            </div>
        </div>

        <!-- 核心指标 -->
        <div class="card">
            <h3>核心指标一览</h3>
            <div class="metrics">
                <div class="metric"><div class="label">初始资金</div><div class="value blue">{INITIAL_CAPITAL:,}</div></div>
                <div class="metric"><div class="label">最终资产</div><div class="value {'green' if m['final']>=INITIAL_CAPITAL else 'red'}">{m['final']:,.0f}</div></div>
                <div class="metric"><div class="label">累计收益率</div><div class="value {'green' if m['total_return']>=0 else 'red'}">{m['total_return']:+.2f}%</div></div>
                <div class="metric"><div class="label">年化收益率</div><div class="value {'green' if m['annual_return']>=0 else 'red'}">{m['annual_return']:+.2f}%</div></div>
                <div class="metric"><div class="label">最大回撤(MDD)</div><div class="value red">{m['mdd']:.2f}%</div></div>
                <div class="metric"><div class="label">夏普比率</div><div class="value {'green' if m['sharpe']>1 else ('blue' if m['sharpe']>0 else 'red')}">{m['sharpe']:.2f}</div></div>
                <div class="metric"><div class="label">Calmar比率</div><div class="value {'green' if m['calmar']>0.5 else 'blue'}">{m['calmar']:.2f}</div></div>
                <div class="metric"><div class="label">总交易次数</div><div class="value blue">{m['trades_total']}</div></div>
                <div class="metric"><div class="label">胜率</div><div class="value {'green' if m['win_rate']>=50 else 'red'}">{m['win_rate']:.1f}%</div></div>
                <div class="metric"><div class="label">平均单笔收益</div><div class="value {'green' if m['avg_trade_return']>=0 else 'red'}">{m['avg_trade_return']:+.2f}%</div></div>
                <div class="metric"><div class="label">年化波动率</div><div class="value blue">{m['volatility']:.2f}%</div></div>
                <div class="metric"><div class="label">持有天数</div><div class="value blue">{m['days']}</div></div>
                <div class="metric"><div class="label">同期买入持有</div><div class="value {'green' if m['buy_hold_return']>=0 else 'red'}">{m['buy_hold_return']:+.2f}%</div></div>
                <div class="metric"><div class="label">超额收益</div><div class="value {'green' if m['excess_return']>=0 else 'red'}">{m['excess_return']:+.2f}%</div></div>
            </div>
        </div>
        <p class="table-title">表_{name} 核心指标汇总</p>

        <!-- 图表 -->
        <div class="card">
            <h3>回测走势图</h3>
            <div class="chart">
                <img src="data:image/png;base64,{m['chart_b64']}" alt="{name}回测走势图">
            </div>
            <p class="chart-title">图_{name}_1 {name}双均线策略回测走势图</p>
            <div class="interp">
                <p>上图展示了 {name} 双均线策略的完整回测过程。上子图为股价走势与均线系统，灰色细线为复权收盘价，红色粗线为MA5短期均线，蓝色粗线为MA15长期均线，红色上行三角标记金叉买入信号，绿色下行三角标记死叉卖出信号，红色区域对应多头市场，蓝色区域对应空头市场。中图为成交量柱状图（阳柱绿色、阴柱红色）。下子图为策略净值曲线，直观反映了策略账户从 {m['first_date']} 到 {m['last_date']} 的价值变化过程。</p>
            </div>
        </div>

        <!-- 买卖信号 -->
        <div class="card">
            <h3>买卖信号汇总</h3>
            <div class="signal-grid">
                <div>
                    <h4 style="color:#e74c3c;">买入信号（{len(buy_dates)}次）</h4>
                    <ul class="signal-list">{buy_detail}</ul>
                </div>
                <div>
                    <h4 style="color:#27ae60;">卖出信号（{len(sell_dates)}次）</h4>
                    <ul class="signal-list">{sell_detail}</ul>
                </div>
            </div>
            <div class="interp">
                <p>在回测期间，{name} 共产生 {len(buy_dates)} 次买入信号和 {len(sell_dates)} 次卖出信号，信号平均间隔约 {signal_interval} 个交易日。{f"信号频率较高，说明该股价格波动较大，均线反复交叉。" if signal_interval < 30 else "信号频率适中。"} 结合股价整体走势，{name} 在此期间呈现<b>{m['trend']}</b>格局，股价从 {m['first_close']:.2f} 元变动至 {m['last_close']:.2f} 元（涨幅 {m['price_change']:+.2f}%）。</p>
            </div>
        </div>

        <!-- 交易记录 -->
        <div class="card">
            <h3>交易记录明细</h3>
            <table>
                <thead><tr><th>序号</th><th>日期</th><th>类型</th><th>价格(元)</th><th>数量(股)</th><th>金额(元)</th><th>剩余现金(元)</th></tr></thead>
                <tbody>{tr_html}</tbody>
            </table>
            <p class="table-title">表_total_{name} 交易记录</p>
            <div class="interp">
                <p>{name} 策略共执行 {len(m['trade_log'])} 笔交易（完整循环 {m['trades_total']//2} 次），胜率 {m['win_rate']:.1f}%（{m['wins']}胜{m['losses']}负），平均单笔收益 {m['avg_trade_return']:+.2f}%。{vs_text}。{sp_text}。回测期间最大回撤为 {m['mdd']:.2f}%，表明策略在最差情况下面临 {abs(m['mdd']):.1f}% 的账户回撤。年化波动率 {m['volatility']:.2f}%，{f"波动较大，持有体验偏剧烈。" if m['volatility'] > 30 else "波动适中。"}</p>
            </div>
        </div>
    </div>
    '''

    stock_js_list.append(stock_html)

# 对比表格行
comparison_rows = ''
for m in results:
    cls_total = 'green' if m['total_return'] >= 0 else 'red'
    cls_sharpe = 'green' if m['sharpe'] > 0 else 'red' if m['sharpe'] < 0 else 'blue'
    cls_win = 'green' if m['win_rate'] >= 50 else 'red'
    cls_excess = 'green' if m['excess_return'] >= 0 else 'red'

    comparison_rows += f'''
    <tr>
        <td><b>{m['name']}</b></td>
        <td class="{cls_total}">{m['total_return']:+.2f}%</td>
        <td class="{cls_total}">{m['annual_return']:+.2f}%</td>
        <td class="red">{m['mdd']:.2f}%</td>
        <td class="{cls_sharpe}">{m['sharpe']:.2f}</td>
        <td>{m['calmar']:.2f}</td>
        <td class="{cls_win}">{m['win_rate']:.1f}%</td>
        <td>{m['trades_total']}</td>
        <td>{m['avg_trade_return']:+.2f}%</td>
        <td>{m['volatility']:.2f}%</td>
        <td class="{cls_excess}">{m['excess_return']:+.2f}%</td>
        <td>{m['trend']}</td>
    </tr>'''

stock_tabs = ''
stock_contents = ''
for i, m in enumerate(results):
    active_tab = ' active' if i == 0 else ''
    active_cont = '' if i > 0 else ''  # 第一个默认显示
    code_key = m['code'].replace('.', '')
    stock_tabs += f'<button class="tab-btn{active_tab}" data-target="{code_key}">{m["name"]}</button>\n'
    stock_contents += stock_js_list[i]

# 三股趋势判断
up_count = sum(1 for m in results if m['total_return'] > 0)
best_stock = max(results, key=lambda x: x['total_return'])
worst_stock = min(results, key=lambda x: x['total_return'])
best_bh = max(results, key=lambda x: x['buy_hold_return'])
worst_bh = min(results, key=lambda x: x['buy_hold_return'])

insight_text = f'''
<p>在 MA{SHORT_WIN}/MA{LONG_WIN} 双均线策略的回测中，三只股票的策略表现呈现明显差异：</p>
<p><b>趋势匹配度</b>：双均线策略的本质是<b>趋势跟踪</b>，其有效性高度依赖于标的资产的走势特征。当股票处于<b>持续单边行情</b>时，均线系统能够较好地跟随趋势方向，金叉死叉信号的有效性较高；但当股票处于<b>震荡行情</b>时，均线频繁交叉，产生大量假信号，策略表现往往不佳。</p>
<p><b>个股差异分析</b>：在本次回测中，{best_stock['name']} 策略表现相对{"" if best_stock['total_return'] < 0 else "较优"}（累计收益 {best_stock['total_return']:+.2f}%），而 {worst_stock['name']} 表现较不理想（累计收益 {worst_stock['total_return']:+.2f}%）。从买入持有收益来看，{best_bh['name']} 涨幅最大（+{best_bh['buy_hold_return']:.2f}%），{worst_bh['name']} 涨幅最小（{worst_bh['buy_hold_return']:+.2f}%）。对比均线策略与买入持有的收益差异，可以判断均线策略是否有效捕捉到了趋势。</p>
<p><b>主要发现</b>：</p>
<ol>
<li>双均线策略在<b>强趋势股</b>中效果更佳——当股价走势持续、方向明确时，金叉死叉信号的准确率显著提高。</li>
<li>在<b>高波动震荡股</b>中，均线策略容易被"左右打脸"，信号频繁但方向不稳定。</li>
<li>短期均线（MA5）对价格变化极其敏感，在震荡市场中易受噪声干扰。<b>适当拉大均线周期差</b>（如MA10/MA30、MA20/MA60）可减少假信号。</li>
<li>双均线策略<b>不适合所有市场环境和所有股票</b>，其本质假设是"过去的价格趋势将持续"，因此在趋势发生逆转的拐点处，策略往往反应滞后，导致回撤。</li>
</ol>
<p><b>适用场景建议</b>：</p>
<ul>
<li><b>推荐</b>：大盘蓝筹股、趋势明确的周期性股票、长期上涨的成长股</li>
<li><b>谨慎</b>：高波动小盘股、横盘震荡的股票、频繁跳空的股票</li>
<li><b>避免</b>：极低流动性的股票、长期处于盘整区的股票</li>
<li><b>改进方向</b>：结合成交量确认、加入ATR波动率过滤器、使用多时间框架确认、采用分批建仓策略</li>
</ul>
'''

# 完整 HTML
html_all = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>三股双均线策略综合分析报告</title>
<style>
* {{ margin:0; padding:0; box-sizing:border-box; }}
body {{ font-family:"SimSun","Times New Roman","Microsoft YaHei",sans-serif; background:#f0f2f5; color:#333; }}
.main-header {{ background:linear-gradient(135deg,#1a1a2e,#16213e,#0f3460); color:#fff; padding:30px 40px; text-align:center; }}
.main-header h1 {{ font-size:24px; letter-spacing:2px; margin-bottom:6px; }}
.main-header .sub {{ font-size:13px; opacity:0.7; }}
.main-header .meta {{ margin-top:10px; display:flex; justify-content:center; flex-wrap:wrap; gap:10px; }}
.main-header .meta span {{ background:rgba(255,255,255,0.12); padding:2px 14px; border-radius:20px; font-size:11px; }}
.tab-bar {{ background:#fff; display:flex; justify-content:center; gap:2px; padding:12px 20px; box-shadow:0 2px 6px rgba(0,0,0,0.06); position:sticky; top:0; z-index:10; }}
.tab-btn {{ padding:10px 28px; font-size:14px; font-family:inherit; font-weight:bold; border:none; border-radius:8px; cursor:pointer; background:#e8edf2; color:#555; transition:all 0.2s; }}
.tab-btn:hover {{ background:#d0d7e0; }}
.tab-btn.active {{ background:#0f3460; color:#fff; }}
.stock-content {{ display:none; max-width:1100px; margin:0 auto; padding:20px; }}
.stock-content:first-of-type {{ display:block; }}
.stock-header {{ background:#fff; border-radius:12px; padding:20px 24px; margin-bottom:18px; box-shadow:0 2px 10px rgba(0,0,0,0.06); }}
.stock-header h2 {{ font-size:20px; color:#1a1a2e; display:flex; align-items:center; gap:10px; }}
.stock-code {{ font-size:12px; color:#888; font-weight:normal; }}
.stock-summary {{ display:flex; flex-wrap:wrap; gap:8px; margin-top:10px; }}
.summary-item {{ font-size:11px; padding:3px 12px; border-radius:14px; background:#f0f2f5; }}
.card {{ background:#fff; border-radius:12px; padding:20px 24px; margin-bottom:16px; box-shadow:0 2px 10px rgba(0,0,0,0.06); }}
.card h3 {{ font-size:15px; color:#1a1a2e; margin-bottom:12px; padding-bottom:6px; border-bottom:2px solid #e8edf2; }}
.metrics {{ display:grid; grid-template-columns:repeat(auto-fit, minmax(150px, 1fr)); gap:10px; }}
.metric {{ background:#f8f9fa; border-radius:8px; padding:12px 8px; text-align:center; }}
.metric .label {{ font-size:10px; color:#888; }}
.metric .value {{ font-size:19px; font-weight:bold; color:#1a1a2e; }}
.value.red {{ color:#e74c3c; }} .value.green {{ color:#27ae60; }} .value.blue {{ color:#2980b9; }}
.chart {{ text-align:center; }}
.chart img {{ max-width:100%; border-radius:6px; }}
.chart-title {{ text-align:center; font-size:10.5pt; font-weight:bold; margin:6px 0 4px; }}
.table-title {{ font-size:10.5pt; font-weight:bold; text-align:center; margin:10px 0 4px; }}
.interp {{ text-align:justify; font-size:10.5pt; line-height:1.5; margin:8px 0; }}
.interp p {{ text-indent:2em; margin-bottom:4px; }}
.signal-grid {{ display:grid; grid-template-columns:1fr 1fr; gap:14px; }}
.signal-list {{ list-style:none; }}
.signal-list li {{ padding:4px 0; border-bottom:1px solid #f0f0f0; font-size:10.5pt; }}
.signal-list li:last-child {{ border:none; }}
table {{ width:100%; border-collapse:collapse; font-size:10.5pt; }}
th {{ background:#1a1a2e; color:#fff; padding:8px 10px; text-align:center; font-weight:500; font-size:10pt; }}
td {{ padding:6px 10px; text-align:center; border-bottom:1px solid #eee; }}
tr:nth-child(even) {{ background:#f8f9fa; }}
.signal-buy {{ color:#e74c3c; font-weight:bold; }}
.signal-sell {{ color:#27ae60; font-weight:bold; }}
.insight-section {{ background:#fff; border-radius:12px; padding:24px 28px; margin:20px auto; max-width:1100px; box-shadow:0 2px 10px rgba(0,0,0,0.06); }}
.insight-section h2 {{ font-size:17px; color:#1a1a2e; margin-bottom:14px; padding-bottom:8px; border-bottom:2px solid #e8edf2; }}
.insight-section p, .insight-section li {{ text-align:justify; font-size:10.5pt; line-height:1.8; text-indent:2em; margin-bottom:6px; }}
.insight-section ol, .insight-section ul {{ padding-left:2em; }}
.insight-section li {{ text-indent:0; }}
.footer {{ text-align:center; font-size:10.5pt; color:#999; padding:20px; }}
@media (max-width:768px) {{ .signal-grid {{ grid-template-columns:1fr; }} .metrics {{ grid-template-columns:repeat(2,1fr); }} }}
</style>
</head>
<body>

<div class="main-header">
<h1>三股双均线策略综合分析报告</h1>
<div class="sub">MA{SHORT_WIN}/{LONG_WIN} 金叉死叉策略 · 金发科技 / 宁德时代 / 比亚迪</div>
<div class="meta">
<span>回测周期: {results[0]['first_date']} ~ {results[0]['last_date']}</span>
<span>初始资金: {INITIAL_CAPITAL:,} 元</span>
<span>数据: Tushare Pro</span>
<span>生成: {datetime.now().strftime("%Y-%m-%d %H:%M")}</span>
</div>
</div>

<div class="tab-bar">
{stock_tabs}
</div>

{stock_contents}

<!-- 综合对比表 -->
<div style="max-width:1100px; margin:20px auto;">
<div class="card">
<h3>三股综合对比（MA{SHORT_WIN}/MA{LONG_WIN}）</h3>
<div style="overflow-x:auto;">
<table>
<thead><tr>
<th>股票</th><th>累计收益</th><th>年化收益</th><th>最大回撤</th><th>夏普</th><th>Calmar</th><th>胜率</th><th>交易次数</th><th>平均单笔</th><th>年化波动</th><th>超额收益</th><th>走势</th>
</tr></thead>
<tbody>
{comparison_rows}
</tbody>
</table>
</div>
<p class="table-title">表_total 三股综合对比表</p>
</div>
</div>

<!-- 总结与心得 -->
<div class="insight-section">
<h2>双均线策略适用场景分析与应用心得</h2>
{insight_text}
</div>

<div class="footer">
<p>报告生成: {datetime.now().strftime("%Y-%m-%d %H:%M")} | 数据来源: Tushare Pro | Python + Matplotlib</p>
</div>

<script>
// Tab 切换
document.querySelectorAll('.tab-btn').forEach(btn => {{
    btn.addEventListener('click', function() {{
        document.querySelectorAll('.tab-btn').forEach(b => b.classList.remove('active'));
        this.classList.add('active');
        document.querySelectorAll('.stock-content').forEach(c => c.style.display = 'none');
        const target = document.getElementById('stock-' + this.dataset.target);
        if (target) target.style.display = 'block';
    }});
}});
</script>

</body>
</html>
'''

html_path = os.path.join(OUT_DIR, '三股双均线策略综合分析报告.html')
with open(html_path, 'w', encoding='utf-8') as f:
    f.write(html_all)
print(f'\nHTML 报告已生成: {html_path} ({os.path.getsize(html_path):,} 字节)')

# ============================================================
# 生成各股 Word 报告
# ============================================================
print('\n生成各股 Word 报告...')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_run_font(run, cn='宋体', en='Times New Roman', sz=Pt(10.5), bold=False):
    run.font.size = sz; run.font.name = en; run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:eastAsia'), cn)
    rpr.rFonts.set(qn('w:ascii'), en)
    rpr.rFonts.set(qn('w:hAnsi'), en)

def add_p(doc, text, indent=Cm(0.74)):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = indent
    run = p.add_run(text)
    set_run_font(run)
    return p

def add_heading_p(doc, text, level=1):
    sz = {0: Pt(16), 1: Pt(14), 2: Pt(12), 3: Pt(10.5)}.get(level, Pt(14))
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT if level <= 2 else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, sz=sz, bold=True)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, sz=Pt(9))
    return p

def add_table_to_doc(doc, data, col_widths=None):
    rows = len(data); cols = len(data[0]) if data else 0
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            set_run_font(run, sz=Pt(9), bold=(i == 0))
    return table

# 各股票Word报告
for m in results:
    doc = Document()
    # 默认样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    name = m['name']; code = m['code']

    # 封面
    doc.add_paragraph(); doc.add_paragraph()
    add_heading_p(doc, f'{name}({code})', 0)
    add_heading_p(doc, f'双均线策略回测报告（MA{SHORT_WIN}/MA{LONG_WIN}）', 2)
    doc.add_paragraph()
    add_caption(doc, f'报告生成: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    add_caption(doc, f'数据来源: Tushare Pro')
    doc.add_page_break()

    # 第一章
    add_heading_p(doc, '一、回测结果概览', 1)
    add_heading_p(doc, '1.1 核心指标', 2)
    metrics_data = [
        ['指标名称', '数值', '说明'],
        ['初始资金', f'{INITIAL_CAPITAL:,}元', '策略启动本金'],
        ['最终资产', f'{m["final"]:,.0f}元', '回测期末总资产'],
        ['累计收益率', f'{m["total_return"]:+.2f}%', '全程总收益率'],
        ['年化收益率', f'{m["annual_return"]:+.2f}%', '折算年化值'],
        ['最大回撤(MDD)', f'{m["mdd"]:.2f}%', '净值最大跌幅'],
        ['夏普比率', f'{m["sharpe"]:.2f}', '风险调整后收益'],
        ['Calmar比率', f'{m["calmar"]:.2f}', '年化收益/最大回撤'],
        ['总交易次数', f'{m["trades_total"]}笔', '回测总交易数'],
        ['胜率', f'{m["win_rate"]:.1f}%', f'{m["wins"]}胜{m["losses"]}负'],
        ['平均单笔收益', f'{m["avg_trade_return"]:+.2f}%', '单笔平均收益率'],
        ['持有天数', f'{m["days"]}天', '总交易日数'],
        ['同期买入持有', f'{m["buy_hold_return"]:+.2f}%', '买入持有收益率'],
        ['超额收益', f'{m["excess_return"]:+.2f}%', '超越买入持有收益'],
    ]
    add_table_to_doc(doc, metrics_data)
    add_caption(doc, f'表1 {name}回测核心指标汇总表')
    add_p(doc, f'表1汇总了{name}本次回测的全部核心指标。累计收益率为{m["total_return"]:+.2f}%，同期买入持有收益为{m["buy_hold_return"]:+.2f}%。年化收益率{m["annual_return"]:+.2f}%，最大回撤{m["mdd"]:.2f}%，夏普比率{m["sharpe"]:.2f}。')

    # 1.2 图表
    add_heading_p(doc, '1.2 回测走势图', 2)
    doc.add_paragraph().add_run().add_picture(m['chart_path'], width=Cm(15))
    add_caption(doc, f'图1 {name}双均线策略回测走势图')
    add_p(doc, f'图1展示了{name}双均线策略的回测过程，由三个子图构成。上图为股价走势与MA5/MA15均线系统，红色上行三角为买入信号（金叉），绿色下行三角为卖出信号（死叉）。中图为成交量柱状图，下图为策略净值曲线，反映账户价值变化。')

    # 1.3 信号
    add_heading_p(doc, '1.3 买卖信号汇总', 2)
    bs = m['buy_signals']; ss = m['sell_signals']
    buy_d = [r['trade_date'] for _, r in bs.iterrows()]
    sell_d = [r['trade_date'] for _, r in ss.iterrows()]
    add_p(doc, f'共产生买入信号{len(bs)}次，卖出信号{len(ss)}次。')
    add_p(doc, f'买入信号日期：{"、".join(buy_d)}')
    add_p(doc, f'卖出信号日期：{"、".join(sell_d)}')
    add_p(doc, f'{name}在回测期间呈现{m["trend"]}格局，股价从{m["first_close"]:.2f}元变动至{m["last_close"]:.2f}元。')

    # 1.4 交易记录
    add_heading_p(doc, '1.4 交易记录明细', 2)
    trade_data = [['序号','日期','类型','价格(元)','金额(元)','剩余现金(元)']]
    for i, t in enumerate(m['trade_log'], 1):
        trade_data.append([str(i), t['date'], t['type'], f'{t["price"]:.2f}', f'{t["amount"]:,.0f}', f'{t["cash"]:,.0f}'])
    add_table_to_doc(doc, trade_data)
    add_caption(doc, f'表2 {name}交易记录明细表')
    add_p(doc, f'{name}策略共执行{len(m["trade_log"])}笔交易，胜率{m["win_rate"]:.1f}%，平均单笔收益{m["avg_trade_return"]:+.2f}%。')

    doc.add_page_break()

    # 第二章
    add_heading_p(doc, '二、量化指标详解', 1)
    add_heading_p(doc, '2.1 累计收益率', 2)
    add_p(doc, f'累计收益率为{m["total_return"]:+.2f}%，初始{INITIAL_CAPITAL:,}元变为{m["final"]:,.0f}元。同期买入持有收益为{m["buy_hold_return"]:+.2f}%。')
    add_heading_p(doc, '2.2 最大回撤', 2)
    add_p(doc, f'最大回撤为{m["mdd"]:.2f}%，意味着净值从最高点最多回撤了{abs(m["mdd"]):.1f}%。')
    add_heading_p(doc, '2.3 夏普比率', 2)
    add_p(doc, f'夏普比率{m["sharpe"]:.2f}，衡量每单位风险获得的超额回报。无风险利率取2.5%。')
    add_heading_p(doc, '2.4 胜率与波动率', 2)
    add_p(doc, f'胜率{m["win_rate"]:.1f}%（{m["wins"]}胜{m["losses"]}负），年化波动率{m["volatility"]:.2f}%。')

    doc.add_page_break()

    # 第三章
    add_heading_p(doc, '三、回测结论', 1)
    add_p(doc, f'在过去一年的回测中，{name}双均线策略(MA{SHORT_WIN}/MA{LONG_WIN})累计收益{m["total_return"]:+.2f}%，同期买入持有{m["buy_hold_return"]:+.2f}%。股票整体呈{m["trend"]}。')
    if m['excess_return'] > 0:
        add_p(doc, f'策略跑赢买入持有{m["excess_return"]:+.2f}%，双均线策略有效捕捉了股价趋势。')
    else:
        add_p(doc, f'策略跑输买入持有{abs(m["excess_return"]):.2f}%，说明在当前参数下策略未能有效捕捉趋势，震荡市场中产生了较多假信号。')

    add_heading_p(doc, '优化建议', 2)
    add_p(doc, '（1）增大均线周期差（如MA10/MA30）以减少假信号。')
    add_p(doc, '（2）引入成交量过滤，金叉时若成交量未放大则暂缓买入。')
    add_p(doc, '（3）结合RSI/MACD辅助指标过滤无效信号。')
    add_p(doc, '（4）采用分批建仓策略，降低单次错误信号影响。')

    word_path = os.path.join(OUT_DIR, f'{name}_双均线策略回测报告.docx')
    doc.save(word_path)
    print(f'  Word报告: {word_path} ({os.path.getsize(word_path):,} 字节)')

print('\n' + '=' * 60)
print('全部完成!')
print('=' * 60)

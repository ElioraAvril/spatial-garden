# -*- coding: utf-8 -*-
"""
金发科技(600143.SH) 双均线策略完整分析
========================================
1) 通过 tushare 获取近一年交易数据并前复权
2) 计算 MA5/MA15 均线
3) 生成金叉买入/死叉卖出信号
4) 可视化：股价+均线+买卖信号+成交量+净值曲线
5) 策略回测与量化指标计算
6) 输出 HTML 报告到 update/
7) 输出 Word 报告到 update/（宋体+TNR，五号字，1.5倍行距，0段间距，两端对齐）
"""

import os, sys, base64, warnings, re, tempfile
warnings.filterwarnings('ignore')
os.environ['MPLCONFIGDIR'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.matplotlib_cache')
from datetime import datetime, timedelta
from io import BytesIO

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

# ===== 字体设置 =====
simsun_path = 'C:/Windows/Fonts/simsun.ttc'
fm.fontManager.addfont(simsun_path)
plt.rcParams['font.family'] = 'SimSun'
plt.rcParams['axes.unicode_minus'] = False

# ===== 输出目录 =====
OUT_DIR = 'D:/项目/北大BA/在线实习/update'
os.makedirs(OUT_DIR, exist_ok=True)

# ===== 参数配置 =====
STOCK_CODE = '600143.SH'
STOCK_NAME = '金发科技'
SHORT_WIN = 5
LONG_WIN = 15
INITIAL_CAPITAL = 1000000
RISK_FREE_RATE = 0.025

print('=' * 60)
print(f'{STOCK_NAME}({STOCK_CODE}) 双均线策略完整分析')
print(f'均线参数: MA{SHORT_WIN} / MA{LONG_WIN}')
print('=' * 60)

# ============================================================
# 1. 获取数据
# ============================================================
print('\n【Step 1】获取数据...')
import tushare as ts
pro = ts.pro_api()

end_date = datetime.now().strftime('%Y%m%d')
start_date = (datetime.now() - timedelta(days=400)).strftime('%Y%m%d')

df = pro.daily(ts_code=STOCK_CODE, start_date=start_date, end_date=end_date)
df.sort_values('trade_date', inplace=True)
df.reset_index(drop=True, inplace=True)
print(f'获取到 {len(df)} 条交易日数据')

# 前复权
df_adj = pro.daily(ts_code=STOCK_CODE, start_date=start_date, end_date=end_date, adj='qfq')
df_adj.sort_values('trade_date', inplace=True)
df_adj.reset_index(drop=True, inplace=True)
df['open_adj'] = df_adj['open']
df['close_adj'] = df_adj['close']
df['high_adj'] = df_adj['high']
df['low_adj'] = df_adj['low']

df['date'] = pd.to_datetime(df['trade_date'])

one_year_ago = df['date'].max() - pd.DateOffset(days=365)
df_plot = df[df['date'] >= one_year_ago].copy()
df_plot.reset_index(drop=True, inplace=True)
print(f'近一年数据: {len(df_plot)} 条')
print(f'日期范围: {df_plot.trade_date.iloc[0]} ~ {df_plot.trade_date.iloc[-1]}')

# ============================================================
# 2. 计算均线
# ============================================================
print('\n【Step 2】计算均线...')
df_plot['MA_short'] = df_plot['close_adj'].rolling(window=SHORT_WIN).mean()
df_plot['MA_long']  = df_plot['close_adj'].rolling(window=LONG_WIN).mean()

# ============================================================
# 3. 买卖信号
# ============================================================
print('\n【Step 3】计算买卖信号...')
df_plot['signal'] = 0
df_plot.loc[(df_plot['MA_short'] > df_plot['MA_long']) &
            (df_plot['MA_short'].shift(1) <= df_plot['MA_long'].shift(1)), 'signal'] = 1
df_plot.loc[(df_plot['MA_short'] < df_plot['MA_long']) &
            (df_plot['MA_short'].shift(1) >= df_plot['MA_long'].shift(1)), 'signal'] = -1

buy_signals  = df_plot[df_plot['signal'] == 1].copy()
sell_signals = df_plot[df_plot['signal'] == -1].copy()
print(f'买入信号: {len(buy_signals)} 次, 卖出信号: {len(sell_signals)} 次')

# ============================================================
# 4. 可视化
# ============================================================
print('\n【Step 4】生成可视化图表...')

fig = plt.figure(figsize=(16, 12), facecolor='white')

# ---- 主图 ----
ax1 = fig.add_axes([0.07, 0.38, 0.88, 0.54])
ax1.set_facecolor('#ffffff')
ax1.plot(df_plot['date'], df_plot['close_adj'], color='#2c3e50', linewidth=1.0, alpha=0.6, label='复权收盘价', zorder=1)
ax1.plot(df_plot['date'], df_plot['MA_short'], color='#e74c3c', linewidth=2.2, label=f'MA{SHORT_WIN} (短期均线)', zorder=2)
ax1.plot(df_plot['date'], df_plot['MA_long'], color='#2980b9', linewidth=2.2, label=f'MA{LONG_WIN} (长期均线)', zorder=2)

# 多空区域
ax1.fill_between(df_plot['date'], df_plot['MA_short'], df_plot['MA_long'],
                 where=(df_plot['MA_short'] >= df_plot['MA_long']),
                 facecolor='#e74c3c', alpha=0.06, zorder=0, label='多头区域')
ax1.fill_between(df_plot['date'], df_plot['MA_short'], df_plot['MA_long'],
                 where=(df_plot['MA_short'] < df_plot['MA_long']),
                 facecolor='#2980b9', alpha=0.06, zorder=0, label='空头区域')

# 买入/卖出标记
ax1.scatter(buy_signals['date'], buy_signals['close_adj'],
            marker='^', s=220, color='#e74c3c', edgecolors='#c0392b', linewidths=1.5, zorder=5, label='买入信号(金叉)')
ax1.scatter(sell_signals['date'], sell_signals['close_adj'],
            marker='v', s=220, color='#27ae60', edgecolors='#1e8449', linewidths=1.5, zorder=5, label='卖出信号(死叉)')

# 标注关键信号
for cnt, (idx, row) in enumerate(buy_signals.iterrows()):
    if cnt == 0 or cnt == len(buy_signals) - 1:
        ax1.annotate(' 买入', (row['date'], row['close_adj']), xytext=(8, 12),
                    textcoords='offset points', fontsize=9, color='#c0392b', fontweight='bold',
                    arrowprops=dict(arrowstyle='->', color='#c0392b', lw=0.8))
for cnt, (idx, row) in enumerate(sell_signals.iterrows()):
    if cnt == 0 or cnt == len(sell_signals) - 1:
        ax1.annotate(' 卖出', (row['date'], row['close_adj']), xytext=(8, -18),
                    textcoords='offset points', fontsize=9, color='#1e8449', fontweight='bold',
                    arrowprops=dict(arrowstyle='->', color='#1e8449', lw=0.8))

ax1.set_title(f'{STOCK_NAME}({STOCK_CODE}) 双均线策略回测  MA{SHORT_WIN}/MA{LONG_WIN}',
              fontsize=16, fontweight='bold', pad=15, color='#2c3e50')
ax1.set_ylabel('价格 (元)', fontsize=12, color='#2c3e50')
ax1.legend(loc='upper left', fontsize=9, framealpha=0.9, edgecolor='#ddd', ncol=2)
ax1.grid(True, alpha=0.15, linestyle='--')
ax1.spines['top'].set_visible(False)
ax1.spines['right'].set_visible(False)

# ---- 成交量 ----
ax2 = fig.add_axes([0.07, 0.24, 0.88, 0.12])
ax2.set_facecolor('#ffffff')
colors_vol = ['#e74c3c' if c < o else '#27ae60' for c, o in zip(df_plot['close_adj'], df_plot['open_adj'])]
ax2.bar(df_plot['date'], df_plot['vol'], color=colors_vol, alpha=0.4, width=0.8, label='成交量')
ax2.set_ylabel('成交量\n(手)', fontsize=10)
ax2.legend(loc='upper left', fontsize=8)
ax2.grid(True, alpha=0.15, linestyle='--')
ax2.spines['top'].set_visible(False)
ax2.spines['right'].set_visible(False)
ax2.set_xticklabels([])

# ---- 净值曲线 ----
position = 0; cash = INITIAL_CAPITAL; nav_series = []
for i in range(len(df_plot)):
    row = df_plot.iloc[i]; price = row['close_adj']
    if row['signal'] == 1 and cash > 0:
        shares = int(cash / price / 100) * 100
        if shares > 0: cash -= shares * price; position += shares
    elif row['signal'] == -1 and position > 0:
        cash += position * price; position = 0
    nav_series.append(cash + position * price)
df_plot['nav'] = nav_series

ax3 = fig.add_axes([0.07, 0.07, 0.88, 0.15])
ax3.set_facecolor('#ffffff')
ax3.plot(df_plot['date'], df_plot['nav'], color='#8e44ad', linewidth=1.8, label='策略净值', zorder=2)
ax3.fill_between(df_plot['date'], INITIAL_CAPITAL, df_plot['nav'],
                 where=(df_plot['nav'] >= INITIAL_CAPITAL), facecolor='#e74c3c', alpha=0.12, label='盈利区域')
ax3.fill_between(df_plot['date'], INITIAL_CAPITAL, df_plot['nav'],
                 where=(df_plot['nav'] < INITIAL_CAPITAL), facecolor='#27ae60', alpha=0.12, label='亏损区域')
ax3.axhline(y=INITIAL_CAPITAL, color='#7f8c8d', linestyle='--', linewidth=0.8, alpha=0.6)
ax3.set_ylabel('净值(万元)', fontsize=10)
ax3.set_xlabel('交易日期', fontsize=11)
ax3.legend(loc='upper left', fontsize=8)
ax3.grid(True, alpha=0.15, linestyle='--')
ax3.spines['top'].set_visible(False)
ax3.spines['right'].set_visible(False)

plt.xticks(rotation=30)

chart_path = os.path.join(OUT_DIR, '金发科技_双均线策略回测图.png')
fig.savefig(chart_path, dpi=200, bbox_inches='tight', facecolor='white')
print(f'图表已保存: {chart_path}')

buf = BytesIO()
fig.savefig(buf, dpi=180, bbox_inches='tight', facecolor='white', edgecolor='none')
buf.seek(0)
chart_b64 = base64.b64encode(buf.read()).decode('utf-8')
plt.close(fig)

# ============================================================
# 5. 策略回测
# ============================================================
print('\n【Step 5】策略回测与量化指标计算...')
position = 0; cash = INITIAL_CAPITAL; trade_log = []

for i in range(len(df_plot)):
    row = df_plot.iloc[i]; price = row['close_adj']
    if row['signal'] == 1 and cash > 0:
        shares = int(cash / price / 100) * 100
        if shares > 0:
            cost = shares * price; cash -= cost; position += shares
            trade_log.append({'date': row['trade_date'], 'type': '买入', 'price': price,
                              'shares': shares, 'amount': cost, 'cash_remain': cash})
    elif row['signal'] == -1 and position > 0:
        amount = position * price; cash += amount; position = 0
        trade_log.append({'date': row['trade_date'], 'type': '卖出', 'price': price,
                          'shares': 0, 'amount': amount, 'cash_remain': cash})

nav_arr = np.array(nav_series)
final_value = nav_arr[-1]
total_return = (final_value - INITIAL_CAPITAL) / INITIAL_CAPITAL * 100
days = (df_plot['date'].iloc[-1] - df_plot['date'].iloc[0]).days
annual_return = (final_value / INITIAL_CAPITAL) ** (365 / max(days, 1)) - 1
annual_return_pct = annual_return * 100

peak_arr = np.maximum.accumulate(nav_arr)
drawdown_arr = (nav_arr - peak_arr) / peak_arr
mdd = drawdown_arr.min() * 100

daily_rets = pd.Series(nav_arr).pct_change().dropna()
daily_mean = daily_rets.mean(); daily_std = daily_rets.std()
annual_mean_r = daily_mean * 252; annual_std_r = daily_std * np.sqrt(252)
sharpe_ratio = (annual_mean_r - RISK_FREE_RATE) / annual_std_r if annual_std_r > 0 else 0

wins = 0; losses = 0; trade_returns = []
for i in range(0, len(trade_log) - 1, 2):
    if i + 1 < len(trade_log) and trade_log[i]['type'] == '买入' and trade_log[i+1]['type'] == '卖出':
        ret = (trade_log[i+1]['amount'] - trade_log[i]['amount']) / trade_log[i]['amount'] * 100
        trade_returns.append(ret)
        if ret > 0: wins += 1
        else: losses += 1
total_trades = wins + losses
win_rate = wins / total_trades * 100 if total_trades > 0 else 0
avg_trade_return = np.mean(trade_returns) if trade_returns else 0

buy_hold_return = (df_plot['close_adj'].iloc[-1] - df_plot['close_adj'].iloc[0]) / df_plot['close_adj'].iloc[0] * 100
excess_return = total_return - buy_hold_return
calmar_ratio = annual_return_pct / abs(mdd) if mdd != 0 else 0

# 输出回测结果
print(f'累计收益率: {total_return:+.2f}% | 年化: {annual_return_pct:+.2f}%')
print(f'最大回撤: {mdd:.2f}% | 夏普比率: {sharpe_ratio:.2f}')
print(f'胜率: {win_rate:.1f}% | 交易: {len(trade_log)}笔')
print(f'买入持有: {buy_hold_return:+.2f}% | 超额: {excess_return:+.2f}%')

# 趋势描述
latest_close = df_plot['close_adj'].iloc[-1]
first_close = df_plot['close_adj'].iloc[0]
price_change_pct = (latest_close - first_close) / first_close * 100
if price_change_pct > 10: trend_desc = '单边上涨趋势'
elif price_change_pct > 0: trend_desc = '震荡偏强走势'
elif price_change_pct > -10: trend_desc = '震荡偏弱走势'
else: trend_desc = '单边下跌趋势'

if excess_return > 0: vs_text = f'策略跑赢买入持有 +{excess_return:.2f}%'
else: vs_text = f'策略跑输买入持有 {excess_return:.2f}%'

if sharpe_ratio > 1: sharpe_comment = f'夏普比率 {sharpe_ratio:.2f}，风险调整后收益较好'
elif sharpe_ratio > 0: sharpe_comment = f'夏普比率 {sharpe_ratio:.2f}，风险调整后收益一般'
else: sharpe_comment = f'夏普比率 {sharpe_ratio:.2f}，风险调整后收益为负，需改进'

# ============================================================
# 6. HTML 报告
# ============================================================
print('\n【Step 6】生成 HTML 报告...')

trade_rows_html = ''
for i, t in enumerate(trade_log, 1):
    cls = 'signal-buy' if t['type'] == '买入' else 'signal-sell'
    trade_rows_html += f'<tr><td>{i}</td><td>{t["date"]}</td><td class="{cls}">{t["type"]}</td><td>{t["price"]:.2f}</td><td>{t["shares"]:,}</td><td>{t["amount"]:,.0f}</td><td>{t["cash_remain"]:,.0f}</td></tr>\n'

html_content = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{STOCK_NAME} 双均线策略回测报告</title>
<style>
* {{ margin:0; padding:0; box-sizing:border-box; }}
body {{ font-family:"SimSun","Times New Roman","Microsoft YaHei",sans-serif; background:#f0f2f5; color:#333; padding:20px; font-size:10.5pt; }}
.container {{ max-width:1100px; margin:0 auto; }}
.header {{ background:linear-gradient(135deg,#1a1a2e,#16213e,#0f3460); color:#fff; padding:35px 40px; border-radius:14px; margin-bottom:24px; }}
.header h1 {{ font-size:26px; margin-bottom:8px; letter-spacing:1px; }}
.header .meta {{ margin-top:10px; display:flex; flex-wrap:wrap; gap:12px; }}
.header .meta span {{ background:rgba(255,255,255,0.15); padding:3px 12px; border-radius:20px; font-size:12px; }}
.card {{ background:#fff; border-radius:12px; padding:24px 28px; margin-bottom:20px; box-shadow:0 2px 12px rgba(0,0,0,0.06); }}
.card h2 {{ font-size:17px; color:#1a1a2e; margin-bottom:14px; padding-bottom:8px; border-bottom:2px solid #e8edf2; }}
.metrics {{ display:grid; grid-template-columns:repeat(auto-fit, minmax(170px, 1fr)); gap:12px; }}
.metric {{ background:#f8f9fa; border-radius:10px; padding:14px 10px; text-align:center; }}
.metric .label {{ font-size:11px; color:#888; margin-bottom:4px; }}
.metric .value {{ font-size:22px; font-weight:bold; color:#1a1a2e; }}
.metric .value.red {{ color:#e74c3c; }}
.metric .value.green {{ color:#27ae60; }}
.metric .value.blue {{ color:#2980b9; }}
.chart {{ text-align:center; margin:10px 0; }}
.chart img {{ max-width:100%; border-radius:8px; box-shadow:0 2px 8px rgba(0,0,0,0.06); }}
.chart-title {{ text-align:center; font-size:10.5pt; font-weight:bold; margin:6px 0 2px 0; }}
.chart-desc {{ text-align:justify; font-size:10.5pt; line-height:1.5; text-indent:2em; margin:6px 12px 12px 12px; color:#555; }}
table {{ width:100%; border-collapse:collapse; font-size:10.5pt; }}
th {{ background:#1a1a2e; color:#fff; padding:10px 12px; text-align:center; font-weight:500; }}
td {{ padding:8px 12px; text-align:center; border-bottom:1px solid #eee; }}
tr:nth-child(even) {{ background:#f8f9fa; }}
.signal-buy {{ color:#e74c3c; font-weight:bold; }}
.signal-sell {{ color:#27ae60; font-weight:bold; }}
.signal-list {{ list-style:none; padding:0; }}
.signal-list li {{ padding:5px 0; border-bottom:1px solid #f0f0f0; font-size:10.5pt; }}
.signal-list li:last-child {{ border:none; }}
.footer {{ text-align:center; font-size:10.5pt; color:#999; padding:20px 0; }}
.conclusion p {{ text-align:justify; font-size:10.5pt; line-height:1.5; text-indent:2em; margin-bottom:6px; }}
.grid-2 {{ display:grid; grid-template-columns:1fr 1fr; gap:14px; }}
.table-title {{ font-size:10.5pt; font-weight:bold; text-align:center; margin:10px 0 4px 0; }}
.section-title {{ font-size:14pt; font-weight:bold; color:#1a1a2e; margin:18px 0 10px 0; padding-bottom:6px; border-bottom:1px solid #ddd; }}
.interp-text {{ text-align:justify; font-size:10.5pt; line-height:1.5; text-indent:2em; margin:4px 0; }}
@media (max-width:768px) {{ .grid-2 {{ grid-template-columns:1fr; }} }}
</style>
</head>
<body>
<div class="container">

<div class="header">
<h1>{STOCK_NAME}({STOCK_CODE}) 双均线策略回测报告</h1>
<div class="meta">
<span>双均线金叉/死叉</span>
<span>MA{SHORT_WIN} / MA{LONG_WIN}</span>
<span>{df_plot["trade_date"].iloc[0]} ~ {df_plot["trade_date"].iloc[-1]}</span>
<span>数据: Tushare Pro</span>
<span>{datetime.now().strftime("%Y-%m-%d %H:%M")}</span>
</div>
</div>

<div class="section-title">一、回测结果概览</div>

<div class="card">
<h2>1.1 核心指标一览</h2>
<div class="metrics">
<div class="metric"><div class="label">初始资金</div><div class="value blue">{INITIAL_CAPITAL:,}</div></div>
<div class="metric"><div class="label">最终资产</div><div class="value {"green" if final_value >= INITIAL_CAPITAL else "red"}">{final_value:,.0f}</div></div>
<div class="metric"><div class="label">累计收益率</div><div class="value {"green" if total_return >= 0 else "red"}">{total_return:+.2f}%</div></div>
<div class="metric"><div class="label">年化收益率</div><div class="value {"green" if annual_return_pct >= 0 else "red"}">{annual_return_pct:+.2f}%</div></div>
<div class="metric"><div class="label">最大回撤(MDD)</div><div class="value red">{mdd:.2f}%</div></div>
<div class="metric"><div class="label">夏普比率</div><div class="value {"green" if sharpe_ratio > 1 else ("blue" if sharpe_ratio > 0 else "red")}">{sharpe_ratio:.2f}</div></div>
<div class="metric"><div class="label">Calmar比率</div><div class="value {"green" if calmar_ratio > 1 else "blue"}">{calmar_ratio:.2f}</div></div>
<div class="metric"><div class="label">交易总次数</div><div class="value blue">{len(trade_log)}</div></div>
<div class="metric"><div class="label">胜率</div><div class="value {"green" if win_rate >= 50 else "red"}">{win_rate:.1f}%</div></div>
<div class="metric"><div class="label">平均单笔收益</div><div class="value {"green" if avg_trade_return >= 0 else "red"}">{avg_trade_return:+.2f}%</div></div>
<div class="metric"><div class="label">年化波动率</div><div class="value blue">{annual_std_r*100:.2f}%</div></div>
<div class="metric"><div class="label">持有天数</div><div class="value blue">{days}</div></div>
<div class="metric"><div class="label">同期买入持有</div><div class="value {"green" if buy_hold_return >= 0 else "red"}">{buy_hold_return:+.2f}%</div></div>
<div class="metric"><div class="label">超额收益</div><div class="value {"green" if excess_return >= 0 else "red"}">{excess_return:+.2f}%</div></div>
</div>
</div>

<p class="table-title">表1 回测核心指标汇总</p>

<div class="card">
<h2>1.2 回测走势图</h2>
<div class="chart">
<img src="data:image/png;base64,{chart_b64}" alt="回测走势图">
</div>
<p class="chart-title">图1 {STOCK_NAME} 双均线策略回测走势图</p>
<p class="chart-desc">
图1展示了金发科技双均线策略的完整回测过程，由三个子图组成。上图为股价走势与均线系统，其中灰色细线为复权收盘价，红色粗线为MA5短期均线，蓝色粗线为MA15长期均线，红色上行三角标记买入信号（金叉），绿色下行三角标记卖出信号（死叉），红色和蓝色半透明区域分别表示多头和空头市场状态，金叉出现在短期均线上穿长期均线的时刻，死叉则相反。中图为每日成交量柱状图，价格收阳时显示为绿色，收阴时显示为红色，成交量放大通常伴随价格重大变动。下图为策略净值曲线，紫色线代表策略账户价值的变化，红色和绿色半透明区域分别表示盈利和亏损区间，灰色水平虚线为初始资金线，底部子图直观反映了策略整体的盈亏过程。
</p>
</div>

<div class="card">
<h2>1.3 交易信号汇总</h2>
<div class="grid-2">
<div>
<h3 style="font-size:10.5pt;color:#e74c3c;margin-bottom:6px;">买入信号（共{len(buy_signals)}次）</h3>
<ul class="signal-list">
{"".join(f'<li>{row["trade_date"]} 价格 {row["close_adj"]:.2f}元 &#x2191; MA{SHORT_WIN}上穿MA{LONG_WIN}</li>' for _, row in buy_signals.iterrows())}
</ul>
</div>
<div>
<h3 style="font-size:10.5pt;color:#27ae60;margin-bottom:6px;">卖出信号（共{len(sell_signals)}次）</h3>
<ul class="signal-list">
{"".join(f'<li>{row["trade_date"]} 价格 {row["close_adj"]:.2f}元 &#x2193; MA{SHORT_WIN}下穿MA{LONG_WIN}</li>' for _, row in sell_signals.iterrows())}
</ul>
</div>
</div>
<p class="interp-text">
从买卖信号分布来看，金发科技在过去一年中均线金叉与死叉交替出现共计{len(buy_signals)+len(sell_signals)}次，平均每约{days//(len(buy_signals)+len(sell_signals)) if len(buy_signals)+len(sell_signals)>0 else 0}个交易日出现一次信号切换，说明该股票价格走势波动较为频繁，均线系统反复交叉，这是震荡市场中双均线策略的典型表现。
</p>
</div>

<div class="card">
<h2>1.4 交易记录明细</h2>
<table>
<thead><tr><th>序号</th><th>日期</th><th>类型</th><th>价格(元)</th><th>数量(股)</th><th>金额(元)</th><th>剩余现金(元)</th></tr></thead>
<tbody>{trade_rows_html}</tbody>
</table>
<p class="table-title">表2 交易记录明细表</p>
<p class="interp-text">
表2列出了全部{len(trade_log)}笔交易的详细记录。策略每次触发金叉时全仓买入（以100股为单位取整），触发死叉时全仓卖出。从交易记录可以看出，每次买入后不久通常会出现卖出信号，多次交易呈现连续亏损，反映了在震荡行情中均线策略频繁切换方向导致的"磨损"效应。{f"值得注意的是，最早一笔买入价格为{trade_log[0]['price']:.2f}元，最后一笔卖出价格为{trade_log[-1]['price']:.2f}元，整体呈现高买低卖的特征。" if len(trade_log) >= 2 else ""}
</p>
</div>

<div class="section-title">二、量化指标详解</div>

<div class="card">
<h2>2.1 累计收益率与年化收益率</h2>
<p class="interp-text">
累计收益率是策略最直观的盈利指标，指从回测开始到结束的总回报率。本次回测累计收益率为<b>{"+" if total_return >=0 else ""}{total_return:.2f}%</b>，年化收益率为<b>{"+" if annual_return_pct >=0 else ""}{annual_return_pct:.2f}%</b>。同期买入持有策略的收益率为<b>{"+" if buy_hold_return >=0 else ""}{buy_hold_return:.2f}%</b>，{vs_text}。
</p>
</div>

<div class="card">
<h2>2.2 最大回撤（Maximum Drawdown）</h2>
<p class="interp-text">
最大回撤（MDD）衡量策略净值从历史最高点回落至后续最低点的最大跌幅，是评价策略风险控制能力的关键指标。本次回测的最大回撤为<b>{mdd:.2f}%</b>，意味着在回测期内，策略净值从最高点最多下跌了{abs(mdd):.1f}%。MDD值越低，代表策略在极端行情下的抗风险能力越强。一般趋势跟踪策略的MDD在20%-40%之间，本次回测的MDD在此范围内。
</p>
</div>

<div class="card">
<h2>2.3 夏普比率（Sharpe Ratio）</h2>
<p class="interp-text">
夏普比率衡量每承担一单位风险所获得的超额回报，计算公式为Sharpe = (Rp - Rf) / σp，其中Rp为年化收益率、Rf为无风险利率（取2.5%）、σp为年化波动率。本次回测夏普比率为<b>{sharpe_ratio:.2f}</b>，{sharpe_comment}。夏普比率越高，说明策略在同等波动水平下创造了更高的收益，收益的"性价比"更好。
</p>
</div>

<div class="card">
<h2>2.4 胜率与平均单笔收益</h2>
<p class="interp-text">
胜率反映策略判断方向的准确程度。本次回测总交易{total_trades}笔（{wins}胜{losses}负），胜率为<b>{win_rate:.1f}%</b>，平均单笔收益为<b>{avg_trade_return:+.2f}%</b>。{f"胜率偏低说明策略在震荡市场中产生了较多假信号，这是双均线策略的主要短板之一。" if win_rate < 50 else "胜率较高，策略判断方向较为准确。"}需要指出的是，即使胜率不高，若盈利交易的平均收益远大于亏损交易，策略仍可能整体盈利——但本次回测中单笔平均收益为负值，说明亏损交易的幅度超过了盈利交易的幅度。
</p>
</div>

<div class="section-title">三、结论与优化建议</div>

<div class="card">
<h2>3.1 回测结论</h2>
<p class="interp-text">
在过去一年的回测中，{STOCK_NAME}双均线策略(MA{SHORT_WIN}/MA{LONG_WIN})累计收益为{total_return:+.2f}%，同期买入持有收益为{buy_hold_return:+.2f}%，{vs_text}。{STOCK_NAME}在过去一年整体呈现<b>{trend_desc}</b>，股价从{first_close:.2f}元变动至{latest_close:.2f}元。{sharpe_comment}，最大回撤{mdd:.2f}%，胜率{win_rate:.1f}%，共产生{len(trade_log)}笔交易。
</p>
</div>

<div class="card">
<h2>3.2 优化建议</h2>
<p class="interp-text">
双均线策略在当前参数(MA{SHORT_WIN}/MA{LONG_WIN})下表现不佳，主要原因是金发科技股价在回测期内经历了明显的单边上涨过程，但均线系统频繁切换方向，未能有效捕捉到这一趋势。据此提出以下优化建议：
</p>
<p class="interp-text">
（1）调整均线参数：增大均线周期差（如MA10/MA30或MA20/MA60），减少假信号频率，但需注意参数越大滞后性越强。
</p>
<p class="interp-text">
（2）成交量过滤：在金叉出现时检查成交量是否明显放大，若成交量未配合则暂缓买入，可以过滤部分无效信号。
</p>
<p class="interp-text">
（3）辅助指标过滤：结合RSI相对强弱指标或MACD指标确认趋势方向，避免在超买区域追涨、超卖区域杀跌。
</p>
<p class="interp-text">
（4）分批建仓：将单次全仓买卖改为分批建仓（如首次金叉建仓50%，回调不破均线再加仓），降低单次错误信号的影响。
</p>
<p class="interp-text">
（5）震荡市识别：加入ATR平均真实波幅或布林带宽度指标，在波动率较低的盘整期间暂停策略。
</p>
</div>

</div>
</body>
</html>
'''

html_path = os.path.join(OUT_DIR, '金发科技_双均线策略回测报告.html')
with open(html_path, 'w', encoding='utf-8') as f:
    f.write(html_content)
print(f'HTML 报告已生成: {html_path} ({os.path.getsize(html_path):,} 字节)')

# ============================================================
# 7. Word 报告
# ============================================================
print('\n【Step 7】生成 Word 报告...')
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ===== 默认样式设置 =====
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)  # 五号
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ===== 辅助函数 =====
def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(10.5), bold=False):
    """设置 run 的中西文字体"""
    run.font.size = size
    run.font.name = en_font
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:eastAsia'), cn_font)
    rpr.rFonts.set(qn('w:ascii'), en_font)
    rpr.rFonts.set(qn('w:hAnsi'), en_font)

def add_paragraph(text, bold=False, size=Pt(10.5), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=Cm(0.74), spacing=1.5):
    """添加带格式的段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = alignment
    if first_indent:
        p.paragraph_format.first_line_indent = first_indent
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p

def add_heading_text(text, level=1):
    """添加标题（宋体加粗，字号按级）"""
    size_map = {0: Pt(16), 1: Pt(14), 2: Pt(12), 3: Pt(10.5)}
    sz = size_map.get(level, Pt(14))
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6) if level > 0 else Pt(12)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT if level <= 2 else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, bold=True, size=sz)
    return p

def add_table_note(text):
    """添加表注（小五号居中）"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=Pt(9))
    return p

def add_chart_image(img_path, caption, width=Cm(14)):
    """插入图片+图注+图注解读"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    # 图注
    add_table_note(caption)

def add_subtitle(text):
    """副标题行（居中，宋体加粗）"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, bold=True, size=Pt(12))
    return p

# ===== 封面 =====
doc.add_paragraph()
doc.add_paragraph()
add_subtitle(f'{STOCK_NAME}({STOCK_CODE})')
add_subtitle('双均线策略回测报告')
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
run = p.add_run(f'MA{SHORT_WIN} / MA{LONG_WIN} 双均线金叉死叉策略')
set_run_font(run)
doc.add_paragraph()
add_table_note(f'报告生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
add_table_note(f'数据来源: Tushare Pro')
doc.add_page_break()

# ===== 第一章：回测结果概览 =====
add_heading_text('一、回测结果概览', level=1)
add_heading_text('1.1 核心指标', level=2)

# 指标表格
metrics_data = [
    ['指标名称', '数值', '说明'],
    ['初始资金', f'{INITIAL_CAPITAL:,} 元', '策略启动时的本金金额'],
    ['最终资产', f'{final_value:,.0f} 元', '回测结束时的账户总资产'],
    ['累计收益率', f'{total_return:+.2f}%', '回测全程的总收益率'],
    ['年化收益率', f'{annual_return_pct:+.2f}%', '将总收益率折算为年化值'],
    ['最大回撤(MDD)', f'{mdd:.2f}%', '净值从峰值到谷值的最大跌幅'],
    ['夏普比率', f'{sharpe_ratio:.2f}', '风险调整后收益指标'],
    ['Calmar比率', f'{calmar_ratio:.2f}', '年化收益与最大回撤之比'],
    ['总交易次数', f'{len(trade_log)} 笔', '回测期间总交易笔数'],
    ['胜率', f'{win_rate:.1f}%', f'盈利交易占比({wins}胜{losses}负)'],
    ['平均单笔收益', f'{avg_trade_return:+.2f}%', '每笔交易的平均收益率'],
    ['持有天数', f'{days} 天', '回测总交易日数'],
    ['同期买入持有', f'{buy_hold_return:+.2f}%', '买入持有策略收益率'],
    ['超额收益', f'{excess_return:+.2f}%', '策略超越买入持有的收益'],
]
table = doc.add_table(rows=len(metrics_data), cols=3, style='Table Grid')
for i, row_data in enumerate(metrics_data):
    for j, cell_text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 2 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(cell_text)
        set_run_font(run, size=Pt(9), bold=(i == 0))

add_table_note('表1 回测核心指标汇总表')
add_table_note('注：无风险利率按2.5%计算，用于夏普比率计算。')

add_paragraph(f'表1汇总了本次回测的全部核心指标。其中累计收益率为{total_return:+.2f}%，同期买入持有收益为{buy_hold_return:+.2f}%，{vs_text}。年化收益率为{annual_return_pct:+.2f}%，最大回撤为{mdd:.2f}%，夏普比率为{sharpe_ratio:.2f}，{sharpe_comment}。')

# ===== 1.2 图表 =====
add_heading_text('1.2 回测走势图', level=2)
add_chart_image(chart_path, '图1 金发科技双均线策略回测走势图', width=Cm(15))

add_paragraph('图1展示了金发科技双均线策略的完整回测过程，由上、中、下三个子图构成。上图为股价走势与均线系统，灰色细线为复权收盘价，红色粗线为MA5短期均线，蓝色粗线为MA15长期均线。从图中可以直观地看到，金发科技在过去一年经历了先涨后跌的走势，股价从约10.5元涨至最高24元附近后回落至约15元。红色上行三角标记买入信号，出现在短期均线上穿长期均线的金叉时刻；绿色下行三角标记卖出信号，出现在短期均线下穿的死叉时刻。红色和蓝色半透明区域分别表示多头和空头市场状态。中图为每日成交量柱状图，下图为策略净值曲线，紫色线直观反映了策略账户价值的变化过程。')

# ===== 1.3 交易信号 =====
add_heading_text('1.3 买卖信号汇总', level=2)
buy_dates_list = [r['trade_date'] for _, r in buy_signals.iterrows()]
sell_dates_list = [r['trade_date'] for _, r in sell_signals.iterrows()]
add_paragraph(f'在回测期间，共产生买入信号 {len(buy_signals)} 次，卖出信号 {len(sell_signals)} 次。')
add_paragraph(f'买入信号（金叉）日期：{", ".join(buy_dates_list)}')
add_paragraph(f'卖出信号（死叉）日期：{", ".join(sell_dates_list)}')
add_paragraph(f'从信号分布来看，金发科技在过去一年中均线金叉与死叉交替出现共计{len(buy_signals)+len(sell_signals)}次，平均每约{days//(len(buy_signals)+len(sell_signals)) if len(buy_signals)+len(sell_signals)>0 else 0}个交易日出现一次信号切换，信号频率较高，说明该股票价格走势波动较为频繁，均线系统反复交叉。这是震荡市场中双均线策略的典型表现——策略在趋势行情中表现较好，但在震荡市中容易产生过多假信号。')

# ===== 1.4 交易记录 =====
add_heading_text('1.4 交易记录明细', level=2)
trade_data = [['序号', '日期', '类型', '价格(元)', '金额(元)', '剩余现金(元)']]
for i, t in enumerate(trade_log, 1):
    trade_data.append([str(i), t['date'], t['type'], f'{t["price"]:.2f}',
                       f'{t["amount"]:,.0f}', f'{t["cash_remain"]:,.0f}'])
table2 = doc.add_table(rows=len(trade_data), cols=6, style='Table Grid')
for i, row_data in enumerate(trade_data):
    for j, cell_text in enumerate(row_data):
        cell = table2.cell(i, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text)
        set_run_font(run, size=Pt(9), bold=(i == 0))

add_table_note('表2 交易记录明细表')
add_paragraph(f'表2列出了全部{len(trade_log)}笔交易的详细记录。策略每次触发金叉时以全部可用资金买入（以100股为单位向下取整），触发死叉时全仓卖出。从交易记录可以看出，多次交易呈现连续亏损特征，反映了在震荡行情中均线策略频繁切换方向导致的"磨损"效应。若最早一笔买入价格为{trade_log[0]["price"]:.2f}元，最后一笔卖出价格为{trade_log[-1]["price"]:.2f}元，整体呈现高买低卖的特征。')

doc.add_page_break()

# ===== 第二章：量化指标详解 =====
add_heading_text('二、量化指标详解', level=1)

add_heading_text('2.1 累计收益率与年化收益率', level=2)
add_paragraph(f'累计收益率是衡量策略盈利能力最直观的指标，指策略从回测开始到结束的总回报率。在本次回测中，策略累计收益率为{total_return:+.2f}%，意味着初始{INITIAL_CAPITAL:,}元的本金在回测期末变为{final_value:,.0f}元。年化收益率将总收益率按持有时间折算为年度收益率，便于不同周期策略间的横向比较，本次年化收益率为{annual_return_pct:+.2f}%。')

add_heading_text('2.2 最大回撤（Maximum Drawdown，MDD）', level=2)
add_paragraph(f'最大回撤是衡量策略风险控制能力的关键指标，它反映的是策略净值从历史最高点回落至后续最低点的最大百分比跌幅。在本次回测中，最大回撤为{mdd:.2f}%，意味着在回测期间，策略净值从最高点最多下跌了{abs(mdd):.1f}%。最大回撤越小，代表策略在极端行情下的抗风险能力越强。在实际投资中，若最大回撤过大（如超过50%），投资者往往难以承受心理压力而在底部割肉离场，无法坚持到后续反转，这就是所谓的"策略死亡"问题。')

add_heading_text('2.3 夏普比率（Sharpe Ratio）', level=2)
add_paragraph(f'夏普比率由诺贝尔经济学奖得主威廉·夏普提出，是衡量风险调整后收益的核心指标。其计算公式为Sharpe = (Rp - Rf) / σp，其中Rp为年化收益率，Rf为无风险利率（取2.5%），σp为年化波动率。本次回测夏普比率为{sharpe_ratio:.2f}，{sharpe_comment}。夏普比率回答的是"为了赚取这些收益，承担了多少风险"这个问题——两个策略收益相同，但波动更小的那个夏普比率更高，收益"质量"更好。')

sharpe_ref = [
    ['夏普比率范围', '评价等级'],
    ['大于2.0', '优秀——风险调整后回报极佳'],
    ['1.0 ~ 2.0', '良好——回报与风险匹配较好'],
    ['0 ~ 1.0', '一般——回报勉强覆盖风险'],
    ['小于0', '不佳——跑输无风险利率'],
]
table3 = doc.add_table(rows=len(sharpe_ref), cols=2, style='Table Grid')
for i, row_data in enumerate(sharpe_ref):
    for j, cell_text in enumerate(row_data):
        cell = table3.cell(i, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text)
        set_run_font(run, size=Pt(9), bold=(i == 0))

add_table_note('表3 夏普比率评估标准')

add_heading_text('2.4 胜率与交易频率', level=2)
add_paragraph(f'本次回测共产生{total_trades}笔完整交易（买入+卖出为一个完整循环），其中盈利交易{wins}笔，亏损交易{losses}笔，胜率为{win_rate:.1f}%，平均单笔收益为{avg_trade_return:+.2f}%。胜率偏低表明在震荡市场中双均线策略的判断准确性不足。值得注意的是，即使胜率不高，若盈利交易的平均盈利幅度远大于亏损交易的平均亏损幅度，策略仍可能实现正收益——但本次回测中单笔平均收益为负值，说明亏损交易的幅度超过了盈利交易。')

doc.add_page_break()

# ===== 第三章：结论与建议 =====
add_heading_text('三、结论与优化建议', level=1)

add_heading_text('3.1 回测结论', level=2)
add_paragraph(f'在过去一年的回测中，金发科技双均线策略(MA{SHORT_WIN}/MA{LONG_WIN})累计收益为{total_return:+.2f}%，同期买入持有收益为{buy_hold_return:+.2f}%，{vs_text}。金发科技在过去一年整体呈现{trend_desc}，股价从{first_close:.2f}元变动至{latest_close:.2f}元（涨幅{price_change_pct:+.2f}%）。{sharpe_comment}，最大回撤{mdd:.2f}%，胜率{win_rate:.1f}%，共产生{len(trade_log)}笔交易。')

add_heading_text('3.2 策略反思', level=2)
add_paragraph(f'双均线策略在当前参数(MA{SHORT_WIN}/MA{LONG_WIN})下表现不佳，核心原因在于金发科技股价在回测期内走势呈现明显波动特征——虽有较大涨幅，但中间反复震荡，均线系统频繁切换方向，多次产生假信号。这说明MA5/MA15的参数组合对该股票而言周期过短，灵敏度过高，容易受到日内波动和短期噪声的干扰。')

add_heading_text('3.3 优化建议', level=2)
add_paragraph('（1）调整均线参数组合。增大均线周期差可以降低信号频率、提高信号可靠性。建议尝试MA10/MA30、MA20/MA60等参数组合，但需要注意参数越大滞后性越强，需要在灵敏度与可靠性之间找到平衡。')
add_paragraph('（2）引入成交量过滤。在金叉出现时，检查当日成交量是否较前几个交易日明显放大。若成交量未能配合，说明市场上涨动力不足，可暂缓买入。')
add_paragraph('（3）结合辅助指标过滤。加入RSI相对强弱指标（避免在超买区域买入）、MACD指标（确认趋势方向）或布林带（识别震荡区间）来减少无效信号。')
add_paragraph('（4）采用分批建仓。将全仓买卖改为分批策略，如首次金叉建仓50%，若价格回调不破长期均线则加仓至100%，降低单次错误信号的影响。')
add_paragraph('（5）在市场盘整期暂停策略。加入ATR平均真实波幅指标，当ATR低于一定阈值时暂停交易，待波动率恢复后再重新启用策略。')

# ===== 保存 =====
word_path = os.path.join(OUT_DIR, '金发科技_双均线策略回测报告.docx')
doc.save(word_path)
print(f'Word 报告已生成: {word_path} ({os.path.getsize(word_path):,} 字节)')

# ===== 完成 =====
print('\n' + '=' * 60)
print('全部完成! 输出文件:')
print(f'  1. HTML 报告:  {html_path}')
print(f'  2. Word 报告:  {word_path}')
print(f'  3. 回测图表:   {chart_path}')
print('=' * 60)
